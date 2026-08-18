"""Generate the binary documents in the demo corpus.

`security_policy.docx` and `employee_handbook.pdf` are committed, because
a demo should not require a build step. This script is how they are
regenerated when their content changes.

    uv run python demo/build_binaries.py

They exist as binaries rather than as more Markdown because two features
depend on the format:

- the DOCX uses real Word `Heading 1/2/3` styles, so structure extraction
  reads the outline from the document rather than guessing it from the
  shape of the text
- the PDF runs to several pages, so citations carry a page number that
  can be checked against the file
"""

from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document
from docx.shared import Pt


DOCUMENTS = Path(__file__).parent / "documents"


# ---------------------------------------------------------------------
# security_policy.docx
# ---------------------------------------------------------------------
#
# (level, text). Level 0 is body text; 1-3 are Heading 1-3.
SECURITY: list[tuple[int, str]] = [
    (1, "Acme Corp Information Security Policy"),
    (0, "Document Type: Information Security Policy"),
    (0, "Document Reference: SEC-POL-2026"),
    (0, "Department: Information Security"),
    (0, "Effective Date: 1 January 2026"),
    (0, "Owner: Chief Information Security Officer"),
    (0, "Classification: Internal"),
    (
        0,
        "This policy states the mandatory security controls that apply "
        "across Acme Corp. Each control has an identifier of the form "
        "SEC-nnnn. Quote the identifier when raising an exception request "
        "or when referring to a control in an audit response.",
    ),
    (
        0,
        "The company is certified to ISO 27001 and is assessed annually. "
        "Controls in this document map to Annex A of that standard; the "
        "mapping is maintained separately by the compliance team.",
    ),
    (1, "1. Access Control"),
    (
        0,
        "Access is granted on the principle of least privilege. An "
        "employee receives the access their role requires and no more. "
        "Access is reviewed when a role changes and removed when "
        "employment ends.",
    ),
    (2, "1.1 Password Requirements"),
    (
        0,
        "Control SEC-4412 blocks the reuse of the previous 12 passwords. "
        "The control is enforced by the identity provider and cannot be "
        "overridden locally.",
    ),
    (
        0,
        "Passwords must be at least 14 characters. Complexity rules are "
        "deliberately not imposed beyond length: length is a better "
        "defence than forced substitution, and complexity rules push "
        "people towards predictable patterns.",
    ),
    (
        0,
        "Passwords are rotated every 180 days for privileged accounts. "
        "Standard accounts are not rotated on a schedule; they are rotated "
        "on evidence of compromise, which is the current guidance from the "
        "National Cyber Security Centre.",
    ),
    (
        0,
        "Every password must be stored in the company password manager. "
        "Storing a password in a browser, a spreadsheet or a note is a "
        "breach of this policy.",
    ),
    (2, "1.2 Multi-Factor Authentication"),
    (
        0,
        "Control SEC-4418 requires multi-factor authentication on every "
        "account that can reach company data. Hardware keys are issued to "
        "privileged users; everyone else uses an authenticator "
        "application.",
    ),
    (
        0,
        "SMS is not an accepted second factor for any account created "
        "after 1 January 2026. Accounts still using SMS are migrated on a "
        "rolling basis.",
    ),
    (2, "1.3 Privileged Access"),
    (
        0,
        "Administrative access is granted just in time and expires after "
        "8 hours. A standing administrative account is permitted only "
        "where a system cannot support just-in-time elevation, and each "
        "such exception is registered and reviewed quarterly.",
    ),
    (
        0,
        "Control SEC-1177 requires that privileged actions are logged to "
        "a store the privileged user cannot modify.",
    ),
    (2, "1.4 Joiners, Movers and Leavers"),
    (
        0,
        "Access is provisioned from the role assigned in the HR system. "
        "A change of role triggers a reassessment within 5 working days. "
        "On the last working day, all access is revoked and devices are "
        "returned.",
    ),
    (
        0,
        "Access reviews are performed quarterly by system owners. A review "
        "that is not completed within 10 working days is escalated to the "
        "department director.",
    ),
    (1, "2. Device Security"),
    (
        0,
        "Employees must use company-approved devices for company work, "
        "and must connect through approved networks or the corporate VPN. "
        "They are expected to lock their computers when away from them.",
    ),
    (2, "2.1 Endpoint Protection"),
    (
        0,
        "Control SEC-2203 requires endpoint detection and response on "
        "every company device. The agent may not be disabled, and a device "
        "reporting the agent as inactive for more than 48 hours is removed "
        "from the network.",
    ),
    (2, "2.2 Encryption"),
    (
        0,
        "Control SEC-2210 requires full-disk encryption on every device "
        "that stores company data, including personal devices covered by "
        "the bring-your-own-device arrangement.",
    ),
    (
        0,
        "Removable media must be encrypted. Unencrypted removable media is "
        "blocked at the endpoint.",
    ),
    (2, "2.3 Patching"),
    (
        0,
        "Operating system updates must be installed within 14 days of "
        "release. Security patches rated critical must be installed within "
        "7 days. Devices more than 30 days behind are blocked from the "
        "network automatically.",
    ),
    (1, "3. Data Classification"),
    (
        0,
        "Every piece of company information carries one of four "
        "classifications. The classification determines how the "
        "information may be stored, shared and destroyed.",
    ),
    (2, "3.1 Public"),
    (
        0,
        "Information approved for release outside the company. Marketing "
        "material, published documentation and the public website. No "
        "handling restrictions.",
    ),
    (2, "3.2 Internal"),
    (
        0,
        "Information intended for employees. Most day-to-day material. May "
        "be shared inside the company freely and outside it only under a "
        "confidentiality agreement.",
    ),
    (2, "3.3 Confidential"),
    (
        0,
        "Information that would cause harm if disclosed. Customer data, "
        "unreleased financial results, source code, contract terms. Access "
        "is granted on a need-to-know basis and every access is logged.",
    ),
    (2, "3.4 Restricted"),
    (
        0,
        "Information that would cause severe harm if disclosed. "
        "Credentials, encryption keys, security assessment findings, and "
        "records of live incidents. Access requires named approval from "
        "the Chief Information Security Officer.",
    ),
    (1, "4. Logging and Monitoring"),
    (
        0,
        "Control SEC-8830 requires that authentication and authorisation "
        "events are logged and retained for 12 months. Network traffic "
        "logs are retained for 90 days.",
    ),
    (
        0,
        "Where this policy and the Data Retention Schedule state different "
        "retention periods for the same record, the longer period applies.",
    ),
    (
        0,
        "Logs are written to an append-only store. No operational account "
        "has permission to delete a log entry.",
    ),
    (1, "5. Incident Response"),
    (
        0,
        "Employees must report suspected security incidents immediately. "
        "Reporting is never penalised, including where the reporter caused "
        "the incident.",
    ),
    (2, "5.1 Reporting"),
    (
        0,
        "Incidents are reported to security@acme.example or to the "
        "24-hour security desk. A suspected phishing message is reported "
        "with the Report Phishing button rather than deleted.",
    ),
    (2, "5.2 Severity"),
    (
        0,
        "Severity 1 is a confirmed breach of Confidential or Restricted "
        "data, or a total loss of a production service. Response begins "
        "within 15 minutes, at any hour.",
    ),
    (
        0,
        "Severity 2 is a suspected breach, or a significant degradation. "
        "Response begins within 1 hour during working hours and within "
        "4 hours outside them.",
    ),
    (
        0,
        "Severity 3 is everything else. Response begins within one working "
        "day.",
    ),
    (2, "5.3 Notification"),
    (
        0,
        "A personal data breach is assessed against the 72-hour regulatory "
        "notification requirement. The Data Protection Officer makes the "
        "notification decision, not the incident responder.",
    ),
    (
        0,
        "Customers are notified where a breach affects their data, "
        "regardless of whether regulatory notification is required.",
    ),
    (2, "5.4 Post-Incident Review"),
    (
        0,
        "Every severity 1 and severity 2 incident has a written review "
        "within 10 working days. Reviews are blameless and are circulated "
        "internally. Control SEC-5540 requires that each review produces "
        "at least one tracked action.",
    ),
    (1, "6. Third Parties"),
    (
        0,
        "A supplier that processes company data is assessed before "
        "engagement and reassessed annually. The assessment covers "
        "certification, subprocessors, breach history and the location of "
        "processing.",
    ),
    (
        0,
        "Control SEC-6612 requires a written data processing agreement "
        "before any Confidential data is shared with a third party.",
    ),
    (1, "7. Physical Security"),
    (
        0,
        "Offices are access-controlled. Badges are personal and must not "
        "be lent. Tailgating — allowing someone to follow you through a "
        "controlled door — is a breach of this policy however polite it "
        "feels.",
    ),
    (
        0,
        "Visitors are signed in, escorted, and issued a visitor badge. "
        "Physical access records are retained for 12 months.",
    ),
    (
        0,
        "A clear desk is required for anything Confidential or above. "
        "Printed material must not be left on a printer.",
    ),
    (1, "8. Exceptions"),
    (
        0,
        "An exception to any control requires a written request naming the "
        "control, the reason, the compensating measure and an expiry date. "
        "Exceptions are approved by the Chief Information Security Officer "
        "and expire after a maximum of 12 months.",
    ),
    (
        0,
        "There is no standing exception to SEC-4412, SEC-4418 or SEC-2210.",
    ),
    (1, "9. Breaches of This Policy"),
    (
        0,
        "A breach may be handled under the disciplinary procedure. "
        "Deliberately disabling a security control, or accessing data "
        "without authorisation, is gross misconduct.",
    ),
    (
        0,
        "Questions about this policy should be directed to "
        "security@acme.example, quoting SEC-POL-2026.",
    ),
]


def build_docx() -> Path:
    """Write the security policy, using real Word heading styles."""

    document = Document()

    # A readable base size; the heading styles carry their own.
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, text in SECURITY:

        if level == 0:
            document.add_paragraph(text)

        else:
            document.add_heading(text, level=level)

    path = DOCUMENTS / "security_policy.docx"
    document.save(path)

    return path


# ---------------------------------------------------------------------
# employee_handbook.pdf
# ---------------------------------------------------------------------
#
# One entry per page, so a citation's page number is checkable by eye.
HANDBOOK: list[tuple[str, str]] = [
    (
        "Acme Corp Employee Handbook",
        """Document Type: Employee Handbook
Document Reference: HR-HBK-2026
Department: People Operations
Effective Date: 1 January 2026
Owner: Director of People Operations

Welcome to Acme Corp.

This handbook is the general guide to working here. It covers how we
expect people to behave, how the everyday processes work, and where to
go when something is wrong.

It is not a contract of employment. Where this handbook and your contract
disagree, your contract applies. Where this handbook and a specific
policy disagree, the specific policy applies — the HR Policy for leave
and notice, the Information Security Policy for security controls, the
Expense Policy for what can be reclaimed.

The handbook is reviewed annually. The version in the People Portal is
always the current one; a printed copy may not be.""",
    ),
    (
        "1. Working at Acme Corp",
        """1.1 Our Expectations

We expect people to be honest, to treat colleagues with respect, and to
do what they said they would do. Most of what follows is an elaboration
of those three things.

1.2 Working Hours

The standard working week is 37.5 hours. Core hours, when everyone is
expected to be contactable, are 10:00 to 16:00. Outside core hours,
arrange your day as suits you and your team.

1.3 Remote Work

Most roles may be worked remotely up to three days a week. Fully remote
arrangements are possible with director approval and a workstation
assessment.

Remote does not mean unavailable. Keep your calendar accurate, and say
where you are working from if it is not your usual place.

1.4 The Office

Offices are open from 07:00 to 20:00 on weekdays. Access outside those
hours requires a badge permission, granted on request.

Desks are unallocated. Book one through the desk booking tool if you are
coming in on a busy day.""",
    ),
    (
        "2. Code of Conduct",
        """2.1 Respect

Harassment, bullying and discrimination are not tolerated. This applies
to conduct in the office, at company events, online, and in any setting
connected to work.

If you experience or witness this, you can raise it informally with your
manager, formally through the grievance procedure in section 5, or
confidentially through the speak-up line.

2.2 Conflicts of Interest

Declare anything that could reasonably be seen as a conflict. Outside
work, a directorship, a family relationship with a supplier, a financial
interest in a competitor. Declaring something is not an admission that
it is a problem; most declared conflicts need no action.

2.3 Gifts and Hospitality

Gifts above 50 GBP may not be accepted. Hospitality may be accepted where
it is proportionate and openly offered. Anything from a public official
requires prior legal review.

The gifts register is maintained by Legal.

2.4 Confidentiality

Company and customer information stays inside the company. This survives
the end of employment.

2.5 Speaking Publicly

Speaking at a conference or writing publicly about your work is
encouraged. Clear the content with your director first if it touches on
customers, unreleased products, or financial information.""",
    ),
    (
        "3. Health, Safety and Wellbeing",
        """3.1 Responsibilities

Acme Corp is responsible for providing a safe place to work. Everyone is
responsible for working safely and for reporting hazards.

3.2 Reporting Accidents

All accidents, however minor, are recorded in the accident book at
reception or through the online form. Near misses should be reported too:
they are the cheapest way to find a hazard.

3.3 First Aid and Fire

First aiders and fire wardens are listed by the lifts on each floor. The
fire alarm is tested on Wednesday mornings at 09:30. If it sounds at any
other time, leave by the nearest exit and go to the assembly point.

3.4 Display Screen Equipment

Everyone working regularly at a screen is entitled to an assessment, in
the office or at home. Eye tests are paid for by the company, along with
a contribution towards glasses where they are needed specifically for
screen work.

3.5 Wellbeing Support

The employee assistance programme is free, confidential, and available 24
hours a day to employees and their immediate family. It covers counselling,
legal and financial guidance.

Mental health first aiders are listed alongside the physical first aiders.""",
    ),
    (
        "4. Disciplinary Procedure",
        """4.1 Principles

The procedure is intended to be fair, consistent and proportionate. Its
purpose is to improve conduct rather than to punish.

You have the right to be accompanied at any formal meeting by a
colleague or a trade union representative.

4.2 Informal Resolution

Most concerns are resolved informally by a conversation with your
manager. An informal conversation is not a disciplinary step and is not
recorded on your file.

4.3 Formal Stages

Stage one is a written warning, live for 6 months.

Stage two is a final written warning, live for 12 months.

Stage three is dismissal, with notice or with pay in lieu.

Stages may be skipped where the conduct is serious enough to warrant it.

4.4 Gross Misconduct

Gross misconduct may result in summary dismissal without notice. Examples
include theft, fraud, violence, deliberate damage, serious breaches of
the Information Security Policy, and falsifying an expense claim.

Where gross misconduct is alleged, you may be suspended on full pay while
it is investigated. Suspension is not a disciplinary sanction and is not
an assumption of guilt.

4.5 Appeals

You may appeal any formal sanction within 10 working days. The appeal is
heard by a manager not previously involved. The appeal decision is final.""",
    ),
    (
        "5. Grievance Procedure",
        """5.1 Informal Stage

An employee should first discuss a grievance informally with their line
manager. Many issues are resolved at this stage, and doing so is not a
lesser outcome.

Where the grievance concerns the line manager, raise it with that
manager's manager or with People Operations instead.

5.2 Formal Stage

If the matter is not resolved informally, submit a formal written
grievance to People Operations, setting out what happened, who was
involved, and what outcome you are seeking.

People Operations must acknowledge the grievance within 5 working days
and hold a hearing within 15 working days.

5.3 The Hearing

You may be accompanied. The hearing is chaired by a manager not
previously involved. You will be given the opportunity to explain the
grievance and to suggest how it might be resolved.

5.4 The Outcome

The outcome is given in writing within 10 working days of the hearing,
with reasons.

5.5 Appeal

You may appeal within 10 working days of receiving the outcome. The
appeal is heard by a more senior manager, and that decision is final.

5.6 Collective Grievances

Where a grievance is raised by several employees about the same matter,
it may be heard as a collective grievance with a nominated
representative.""",
    ),
    (
        "6. Probation",
        """6.1 The Probation Period

Probation terms are set out in the HR Policy and are not repeated here.
In summary, new employees serve a probation period, reviews are held part
way through and at the end, and the outcome is confirmed in writing.

6.2 What Probation Affects

While on probation you are not eligible for the discretionary bonus
scheme, the training budget, or unpaid leave.

International travel during probation requires additional approval, as
set out in the Travel Booking Policy.

Your notice period during probation is shorter than the standard notice
period; the exact terms are in the HR Policy.

6.3 Support During Probation

Probation is a two-way assessment. You should expect regular one-to-ones,
a clear plan for the first three months, and a buddy who is not your
manager.

If something is not working, say so early. A probation period that ends
badly is almost always one where a problem was visible for weeks and not
discussed.

6.4 After Probation

On successful completion you are enrolled in the private medical scheme,
the standard notice period applies, and the full benefits package becomes
available.""",
    ),
    (
        "7. Leaving Acme Corp",
        """7.1 Resignation

Resignation must be in writing to your line manager and to People
Operations. The notice period is set out in the HR Policy and in your
contract.

7.2 Handover

You are expected to produce a written handover covering current work,
outstanding commitments, key contacts, and anything only you know. Your
manager will agree the scope with you at the start of the notice period.

7.3 Garden Leave

The company may ask you not to attend work during your notice while
remaining employed and paid. Confidentiality obligations continue
throughout.

7.4 Return of Property

Laptops, badges, phones, keys and any printed material are returned on
the last working day.

7.5 Final Pay

Final pay includes salary to the last day and payment for accrued but
untaken annual leave. Leave taken in excess of the accrued entitlement is
recovered from the final payment.

7.6 References

Acme Corp provides factual references confirming job title and dates of
employment. Requests go to People Operations, not to individual managers.

7.7 Exit Interview

An exit interview is offered to everyone who leaves. It is voluntary, and
what is said is reported in aggregate rather than attributed.

7.8 Staying in Touch

The alumni network is open to anyone who leaves on good terms. Former
employees are eligible to be rehired and are not disadvantaged by having
left.""",
    ),
]


def build_pdf() -> Path:
    """Write the handbook, one section per page."""

    document = pymupdf.open()

    for index, (title, body) in enumerate(HANDBOOK, start=1):

        page = document.new_page()

        page.insert_text(
            (72, 84),
            title,
            fontsize=16,
            fontname="helv",
        )

        # A text box rather than insert_text, so long paragraphs wrap
        # instead of running off the page.
        page.insert_textbox(
            pymupdf.Rect(72, 110, 523, 740),
            body,
            fontsize=10.5,
            fontname="helv",
            align=0,
        )

        page.insert_text(
            (72, 760),
            f"Acme Corp Employee Handbook — page {index}",
            fontsize=8,
            fontname="helv",
        )

    path = DOCUMENTS / "employee_handbook.pdf"
    document.save(path)
    document.close()

    return path


if __name__ == "__main__":

    for path in (build_docx(), build_pdf()):
        print(f"  wrote {path.relative_to(Path.cwd())} ({path.stat().st_size:,} bytes)")
