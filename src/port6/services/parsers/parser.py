"""Document parsing.

Parsers emit an ordered list of blocks rather than one flat string, so that
page numbers and heading structure survive into retrieval. `ParsedDocument.text`
is still the joined plain text, which keeps existing callers working.

Where a text layer does not cover the file, OCR fills the gap — see
`ocr.py` for how a page is classified and why it is done per page. OCR runs
inside the upload request, so it is bounded by `ocr.max_pages`: a file
needing more than that is refused up front rather than left to time out.
"""

from pathlib import Path
import re

import markdown
from pydantic import BaseModel, Field
import pymupdf
from bs4 import BeautifulSoup
from docx import Document

from port6.services.parsers import ocr
from port6.services.settings.service import get_int, get_setting


# Matches "4.", "4.2", "4.2.1" and "Section 4" style numbering.
NUMBERED_HEADING = re.compile(
    r"^(?:section\s+)?(\d+(?:\.\d+)*)\.?\s+(\S.*)$",
    re.IGNORECASE,
)

# A heading is short; anything longer is almost certainly a sentence.
MAX_HEADING_CHARACTERS = 120

# "Version: 6.2" — a labelled metadata field rather than a section heading.
LABELLED_LINE = re.compile(r"^[A-Za-z][A-Za-z ]{1,30}:\s*\S")


class ParsedBlock(BaseModel):
    """One paragraph or heading, with where it came from."""

    text: str
    page_number: int | None = None

    # None for body text; 1..6 for a heading, where 1 is the top level.
    heading_level: int | None = None

    @property
    def is_heading(self) -> bool:
        return self.heading_level is not None


class OcrLimitExceeded(ValueError):
    """More pages need OCR than the configured budget allows.

    Raised before any OCR runs, so the cost of refusing is the cheap
    classification pass and nothing more.
    """


class ParsedDocument(BaseModel):
    source: str
    file_type: str
    text: str
    blocks: list[ParsedBlock] = Field(default_factory=list)

    # Which pages OCR actually contributed text to, 1-based. Empty for a
    # document whose text layer covered it, which is the common case.
    ocr_pages: list[int] = Field(default_factory=list)


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = [line.strip() for line in text.splitlines()]

    cleaned_lines = []
    previous_empty = False

    for line in lines:
        if not line:
            if not previous_empty:
                cleaned_lines.append("")
            previous_empty = True
        else:
            cleaned_lines.append(line)
            previous_empty = False

    return "\n".join(cleaned_lines).strip()


def detect_heading_level(
    line: str,
) -> int | None:
    """Guess whether a line is a heading, for formats that do not say so.

    Used for PDF and plain text, where there is no style information. DOCX
    and Markdown carry explicit levels and do not need this.
    """

    stripped = line.strip()

    if not stripped or len(stripped) > MAX_HEADING_CHARACTERS:
        return None

    # Sentences and list items are not headings.
    if stripped.endswith((".", ",", ";", ":")) and not NUMBERED_HEADING.match(
        stripped
    ):
        return None

    match = NUMBERED_HEADING.match(stripped)

    if match:
        # A numbered heading sits under the document title, so "1." is
        # level 2 and "1.1" is level 3 — mirroring how the same document
        # would be written in Markdown (# title, ## 1., ### 1.1).
        depth = len(match.group(1).split("."))
        return min(depth + 1, 6)

    # "Version: 6.2" and "Effective From: 2026-01-01" are metadata fields,
    # not sections. Without this a document's header block would become a
    # handful of one-line sections.
    if LABELLED_LINE.match(stripped):
        return None

    words = stripped.split()

    # ALL CAPS lines of a few words read as section headers.
    if (
        stripped.isupper()
        and 1 <= len(words) <= 12
        and any(character.isalpha() for character in stripped)
    ):
        return 1

    # Title Case with no terminal punctuation, e.g. "Maternity Leave".
    if (
        1 < len(words) <= 8
        and not stripped.endswith((".", "?", "!"))
        and all(
            word[0].isupper()
            for word in words
            if word and word[0].isalpha()
        )
    ):
        return 1

    return None


def blocks_to_text(
    blocks: list[ParsedBlock],
) -> str:
    return clean_text(
        "\n\n".join(block.text for block in blocks)
    )


def _blocks_from_lines(
    text: str,
    page_number: int | None = None,
) -> list[ParsedBlock]:
    """Split plain text into blocks, inferring headings by shape."""

    blocks: list[ParsedBlock] = []
    paragraph: list[str] = []

    def flush() -> None:
        if not paragraph:
            return

        joined = "\n".join(paragraph).strip()

        if joined:
            blocks.append(
                ParsedBlock(
                    text=joined,
                    page_number=page_number,
                )
            )

        paragraph.clear()

    for line in text.splitlines():

        stripped = line.strip()

        if not stripped:
            flush()
            continue

        level = detect_heading_level(stripped)

        if level is not None:
            flush()
            blocks.append(
                ParsedBlock(
                    text=stripped,
                    page_number=page_number,
                    heading_level=level,
                )
            )
            continue

        paragraph.append(stripped)

    flush()

    return blocks


def parse(path: Path) -> ParsedDocument:
    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if not path.is_file():
        raise ValueError(f"Path is not a file: {path}")

    extension = path.suffix.lower()

    parsers = {
        ".pdf": pdf_parser,
        ".docx": doc_parser,
        ".txt": txt_parser,
        ".md": md_parser,
    }

    if extension not in parsers:
        raise ValueError(f"Unsupported file type: {extension}")

    return parsers[extension](path)


def _ocr_budget() -> tuple[bool, int, int]:
    """`(enabled, dpi, max_units)` for this parse.

    Read once per document rather than per page: these are database-backed
    settings, and a page loop is the wrong place to be asking.
    """

    if not get_setting("ocr.enabled") or not ocr.is_available():
        return False, 0, 0

    return True, get_int("ocr.dpi"), get_int("ocr.max_pages")


def _ocr_modes(document) -> dict[int, str]:
    """Which pages to OCR, refusing up front if there are too many.

    The budget is checked against the whole plan before a single page is
    rasterised, which is what makes the refusal cheap — the classification
    pass is text lengths and image counts, measured at 9ms for 8 pages.
    """

    enabled, _, max_units = _ocr_budget()

    if not enabled:
        return {}

    modes = ocr.pages_to_ocr(document)

    if len(modes) > max_units:
        raise OcrLimitExceeded(
            f"this file needs OCR on {len(modes)} of its "
            f"{len(document)} pages, over the limit of {max_units}. "
            "Raise the ocr.max_pages setting to accept files this large, "
            "or split it into smaller uploads."
        )

    return modes


def _no_pdf_text_message(path: Path) -> str:
    """Why a PDF yielded nothing, in terms of what to do about it."""

    if not get_setting("ocr.enabled"):
        return (
            f"No text found in PDF: {path}. It looks scanned, and OCR is "
            "switched off — enable the ocr.enabled setting to read it."
        )

    if not ocr.is_available():
        return (
            f"No text found in PDF: {path}. It looks scanned, and OCR is "
            "unavailable — install the tesseract binary to read it."
        )

    return f"No text found in PDF: {path}"


def pdf_parser(path: Path) -> ParsedDocument:
    blocks: list[ParsedBlock] = []
    ocr_pages: list[int] = []

    with pymupdf.open(path) as doc:

        modes = _ocr_modes(doc)
        _, dpi, _ = _ocr_budget()

        for page_index, page in enumerate(doc, start=1):

            text = clean_text(page.get_text())

            mode = modes.get(page_index - 1)

            if mode is not None:
                recovered = clean_text(
                    ocr.ocr_page(
                        page,
                        dpi=dpi,
                        full=mode == ocr.OCR_FULL,
                    )
                )

                # Only taken when it actually added something. On a page
                # whose images hold no text, OCR returns the text layer
                # back unchanged, and a failed page returns less than it —
                # neither is a reason to discard what was already read.
                if len(recovered) > len(text):
                    text = recovered
                    ocr_pages.append(page_index)

            if not text:
                continue

            # Page numbers are only available here, so they are attached
            # before the pages get joined together.
            blocks.extend(
                _blocks_from_lines(
                    text,
                    page_number=page_index,
                )
            )

    if not blocks:
        raise ValueError(_no_pdf_text_message(path))

    return ParsedDocument(
        source=path.name,
        file_type="pdf",
        text=blocks_to_text(blocks),
        blocks=blocks,
        ocr_pages=ocr_pages,
    )


def _docx_body_images(doc) -> list:
    """Image parts the document body actually references.

    `related_parts` rather than every part in the package: Word stores a
    preview thumbnail under docProps, which is not content and OCRs to
    nothing. The one .docx in this project has exactly that and no body
    images, so enumerating the package would have found a thumbnail and
    called it a figure.
    """

    return [
        part
        for part in doc.part.related_parts.values()
        if part.content_type.startswith("image/")
    ]


def _docx_image_blocks(doc) -> list[ParsedBlock]:
    """Text recovered from figures pasted into a DOCX.

    A screenshot of a table is the common case, and without this its
    numbers are in the file but invisible to retrieval. Labelled like the
    tables above so a reader can tell where the text came from; no page
    number, because DOCX has no pages.
    """

    images = _docx_body_images(doc)

    if not images:
        return []

    enabled, dpi, max_units = _ocr_budget()

    if not enabled:
        return []

    # The same budget as PDF pages, counted in images: it bounds how long
    # one upload can spend in OCR, whatever the unit happens to be.
    if len(images) > max_units:
        raise OcrLimitExceeded(
            f"this file has {len(images)} images to OCR, over the limit "
            f"of {max_units}. Raise the ocr.max_pages setting to accept "
            "files this large."
        )

    blocks = []

    for number, part in enumerate(images, start=1):

        text = clean_text(ocr.ocr_image_bytes(part.blob, dpi=dpi))

        if not text:
            continue

        blocks.append(
            ParsedBlock(
                text=f"Image {number}:\n{text}",
            )
        )

    return blocks


def doc_parser(path: Path) -> ParsedDocument:
    doc = Document(path)
    blocks: list[ParsedBlock] = []

    for paragraph in doc.paragraphs:

        text = clean_text(paragraph.text)

        if not text:
            continue

        # python-docx exposes the real outline, so no guessing is needed.
        style_name = (paragraph.style.name or "") if paragraph.style else ""

        heading_level = None

        if style_name.lower().startswith("heading"):
            digits = "".join(
                character
                for character in style_name
                if character.isdigit()
            )
            heading_level = int(digits) if digits else 1

        elif style_name.lower() == "title":
            heading_level = 1

        else:
            # Some documents mark headings by formatting alone.
            heading_level = detect_heading_level(text)

        blocks.append(
            ParsedBlock(
                text=text,
                heading_level=heading_level,
            )
        )

    for table_number, table in enumerate(doc.tables, start=1):

        table_lines = [f"Table {table_number}:"]

        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            table_lines.append(" | ".join(cells))

        blocks.append(
            ParsedBlock(
                text="\n".join(table_lines),
            )
        )

    blocks.extend(_docx_image_blocks(doc))

    if not blocks:
        raise ValueError(f"No text found in DOCX: {path}")

    return ParsedDocument(
        source=path.name,
        file_type="docx",
        text=blocks_to_text(blocks),
        blocks=blocks,
    )


def txt_parser(path: Path) -> ParsedDocument:
    text = clean_text(get_text_from_file(path))

    if not text:
        raise ValueError(f"TXT file is empty: {path}")

    blocks = _blocks_from_lines(text)

    return ParsedDocument(
        source=path.name,
        file_type="txt",
        text=text,
        blocks=blocks,
    )


def md_parser(path: Path) -> ParsedDocument:
    raw = get_text_from_file(path)

    if not raw.strip():
        raise ValueError(f"Markdown file is empty: {path}")

    html = markdown.markdown(
        raw,
        extensions=[
            "tables",
            "fenced_code",
        ],
    )

    soup = BeautifulSoup(html, "html.parser")

    blocks: list[ParsedBlock] = []

    # Markdown states its heading levels outright, so walk the rendered
    # elements rather than re-inferring structure from the text.
    for element in soup.find_all(
        [
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "p",
            "li",
            "pre",
            "table",
            "blockquote",
        ]
    ):
        # Skip nodes whose text is already covered by an ancestor block.
        if element.find_parent(
            ["li", "pre", "table", "blockquote"]
        ):
            continue

        text = clean_text(element.get_text("\n"))

        if not text:
            continue

        heading_level = (
            int(element.name[1])
            if element.name.startswith("h") and element.name[1:].isdigit()
            else None
        )

        blocks.append(
            ParsedBlock(
                text=text,
                heading_level=heading_level,
            )
        )

    if not blocks:
        raise ValueError(f"No text found in Markdown: {path}")

    return ParsedDocument(
        source=path.name,
        file_type="md",
        text=blocks_to_text(blocks),
        blocks=blocks,
    )


def get_text_from_file(path: Path) -> str:
    with open(path, "r", encoding="utf-8") as file:
        return file.read()
