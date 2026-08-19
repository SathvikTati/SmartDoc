"""OCR: the text a text layer does not carry.

A scanned PDF is a picture of a page. Before OCR the parser read nothing
from it and rejected the file as empty, so these fix the three cases a real
library mixes together — a digital page, a scanned page, and a digital page
with a scanned figure on it.

The fixtures are built rather than committed: a binary scan in the repo
could not be reviewed, and building one states exactly what makes it
scanned — a raster with no text layer behind it.
"""

from __future__ import annotations

import pymupdf
import pytest
from docx import Document as DocxDocument

from port6.services.parsers import ocr
from port6.services.parsers import parser
from port6.services.parsers.parser import OcrLimitExceeded, parse


needs_tesseract = pytest.mark.skipif(
    not ocr.is_available(),
    reason="the tesseract binary is not installed",
)


# --- fixtures ---------------------------------------------------------

def _raster(lines: list[str], width: int = 380, height: int = 220):
    """Text rendered to pixels, at a size that stays legible.

    Legibility is the whole point of the size argument. A raster squeezed
    into a small box on the page downscales into noise and OCRs to nothing
    — which looks exactly like a bug in the parser.
    """

    doc = pymupdf.open()
    page = doc.new_page(width=width, height=height)

    y = 40

    for line in lines:
        page.insert_text((20, y), line, fontsize=13)
        y += 26

    pixmap = page.get_pixmap(dpi=150)
    doc.close()

    return pixmap


def digital_pdf(path, lines: list[str]) -> str:
    doc = pymupdf.open()
    page = doc.new_page()

    y = 80

    for line in lines:
        page.insert_text((50, y), line, fontsize=12)
        y += 24

    doc.save(path)
    doc.close()

    return str(path)


def scanned_pdf(path, lines: list[str], pages: int = 1) -> str:
    """A PDF whose pages are images and nothing else."""

    pixmap = _raster(lines)

    doc = pymupdf.open()

    for _ in range(pages):
        page = doc.new_page()
        page.insert_image(page.rect, pixmap=pixmap)

    doc.save(path)
    doc.close()

    return str(path)


def mixed_pdf(path, text_lines: list[str], image_lines: list[str]) -> str:
    """A real text layer, and a figure holding text of its own."""

    pixmap = _raster(image_lines)

    doc = pymupdf.open()
    page = doc.new_page()

    y = 80

    for line in text_lines:
        page.insert_text((50, y), line, fontsize=12)
        y += 24

    page.insert_image(
        pymupdf.Rect(50, y + 10, 50 + pixmap.width, y + 10 + pixmap.height),
        pixmap=pixmap,
    )

    doc.save(path)
    doc.close()

    return str(path)


# --- classification ---------------------------------------------------

class TestPageClassification:
    """Per page, because one file mixes the cases."""

    def test_a_digital_page_is_left_alone(self, tmp_path):
        path = digital_pdf(tmp_path / "digital.pdf", ["3. Notice Periods"])

        with pymupdf.open(path) as doc:
            assert ocr.pages_to_ocr(doc) == {}

    def test_a_scanned_page_is_read_whole(self, tmp_path):
        path = scanned_pdf(tmp_path / "scan.pdf", ["SCANNED HEADING"])

        with pymupdf.open(path) as doc:
            assert ocr.pages_to_ocr(doc) == {0: ocr.OCR_FULL}

    def test_a_page_with_both_has_only_its_images_read(self, tmp_path):
        path = mixed_pdf(
            tmp_path / "mixed.pdf",
            ["3. Annual Leave"],
            ["GRADE A 22 DAYS"],
        )

        with pymupdf.open(path) as doc:
            assert ocr.pages_to_ocr(doc) == {0: ocr.OCR_IMAGES}

    def test_a_stray_character_does_not_count_as_a_text_layer(self, tmp_path):
        """Scanners stamp things on. A few characters is not a text layer."""

        pixmap = _raster(["SCANNED HEADING"])

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_image(page.rect, pixmap=pixmap)
        page.insert_text((10, 10), "1", fontsize=6)

        path = str(tmp_path / "stamped.pdf")
        doc.save(path)
        doc.close()

        with pymupdf.open(path) as opened:
            assert ocr.pages_to_ocr(opened) == {0: ocr.OCR_FULL}


# --- parsing ----------------------------------------------------------

@needs_tesseract
class TestScannedPdf:

    def test_a_scanned_pdf_is_read_instead_of_refused(self, tmp_path):
        path = scanned_pdf(
            tmp_path / "policy.pdf",
            ["ACME CORP HR POLICY", "28 WEEKS OF LEAVE"],
        )

        parsed = parse(path)

        assert "ACME" in parsed.text.upper()
        assert parsed.ocr_pages == [1]

    def test_the_page_a_line_came_from_survives_ocr(self, tmp_path):
        path = scanned_pdf(tmp_path / "two.pdf", ["SCANNED HEADING"], pages=2)

        parsed = parse(path)

        assert parsed.ocr_pages == [1, 2]
        assert {block.page_number for block in parsed.blocks} == {1, 2}


@needs_tesseract
class TestMixedPdf:

    def test_text_in_a_figure_reaches_the_text(self, tmp_path):
        path = mixed_pdf(
            tmp_path / "mixed.pdf",
            ["3. Annual Leave Entitlement"],
            ["GRADE A 22 DAYS"],
        )

        parsed = parse(path)

        assert "Annual Leave Entitlement" in parsed.text
        assert "22 DAYS" in parsed.text.upper()

    def test_the_text_layer_is_not_read_twice(self, tmp_path):
        """`full=False` is what makes this safe; assert it stays that way."""

        path = mixed_pdf(
            tmp_path / "mixed.pdf",
            ["3. Annual Leave Entitlement"],
            ["GRADE A 22 DAYS"],
        )

        parsed = parse(path)

        assert parsed.text.count("Annual Leave Entitlement") == 1


@needs_tesseract
class TestDigitalPdfIsUnaffected:
    """The common case must pay nothing and read identically."""

    def test_a_digital_pdf_parses_the_same_with_ocr_available(self, tmp_path):
        path = digital_pdf(
            tmp_path / "digital.pdf",
            ["3. Notice Periods", "Employees give 60 days written notice."],
        )

        parsed = parse(path)

        assert parsed.ocr_pages == []
        assert "60 days written notice" in parsed.text


# --- the budget -------------------------------------------------------

class TestPageBudget:
    """OCR runs inside the upload request, so it has to be bounded."""

    def test_too_many_pages_is_refused_before_any_are_read(
        self,
        tmp_path,
        monkeypatch,
    ):
        monkeypatch.setattr(
            parser,
            "get_int",
            lambda key: 3 if key == "ocr.max_pages" else 200,
        )

        path = scanned_pdf(tmp_path / "big.pdf", ["SCANNED"], pages=5)

        with pytest.raises(OcrLimitExceeded) as caught:
            parse(path)

        # The message has to name both numbers, or there is nothing to act on.
        assert "5" in str(caught.value)
        assert "3" in str(caught.value)

    @needs_tesseract
    def test_a_file_inside_the_budget_is_read(self, tmp_path, monkeypatch):
        monkeypatch.setattr(
            parser,
            "get_int",
            lambda key: 5 if key == "ocr.max_pages" else 200,
        )

        path = scanned_pdf(tmp_path / "ok.pdf", ["SCANNED HEADING"], pages=3)

        assert len(parse(path).ocr_pages) == 3


# --- switched off, or unavailable ------------------------------------

class TestWithoutOcr:

    def test_a_scanned_pdf_says_why_when_ocr_is_off(self, tmp_path, monkeypatch):
        monkeypatch.setattr(
            parser,
            "get_setting",
            lambda key: False if key == "ocr.enabled" else True,
        )

        path = scanned_pdf(tmp_path / "scan.pdf", ["SCANNED HEADING"])

        with pytest.raises(ValueError, match="switched off"):
            parse(path)

    def test_a_digital_pdf_still_parses_with_ocr_off(self, tmp_path, monkeypatch):
        monkeypatch.setattr(
            parser,
            "get_setting",
            lambda key: False if key == "ocr.enabled" else True,
        )

        path = digital_pdf(tmp_path / "digital.pdf", ["3. Notice Periods"])

        assert "Notice Periods" in parse(path).text

    def test_a_missing_binary_is_reported_as_such(self, tmp_path, monkeypatch):
        monkeypatch.setattr(ocr, "is_available", lambda: False)

        path = scanned_pdf(tmp_path / "scan.pdf", ["SCANNED HEADING"])

        with pytest.raises(ValueError, match="unavailable"):
            parse(path)


# --- DOCX -------------------------------------------------------------

@needs_tesseract
class TestDocxImages:

    def test_a_figure_pasted_into_a_docx_is_read(self, tmp_path):
        image = tmp_path / "table.png"
        _raster(["TABLE 7: EXPENSE LIMITS", "MEALS 35 GBP"]).save(image)

        document = DocxDocument()
        document.add_paragraph("4. Expenses")
        document.add_picture(str(image))

        path = tmp_path / "expenses.docx"
        document.save(path)

        parsed = parse(path)

        assert "4. Expenses" in parsed.text
        assert "EXPENSE LIMITS" in parsed.text.upper()

    def test_the_preview_thumbnail_is_not_mistaken_for_a_figure(self, tmp_path):
        """Word keeps a thumbnail under docProps. It is not content.

        Enumerating every image part in the package would find it, OCR it to
        nothing, and add an empty figure block to every document Word saved.
        """

        document = DocxDocument()
        document.add_paragraph("1. Scope")

        path = tmp_path / "plain.docx"
        document.save(path)

        opened = DocxDocument(path)

        assert parser._docx_body_images(opened) == []
